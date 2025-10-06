from __future__ import annotations

import io
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from bs4 import BeautifulSoup
from PIL import Image
import io as _io
import pytesseract
import csv as _csv
import json as _json
from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document as DocxDocument

from ..schemas import ImageRef, ParsedBundle, TableRef
from ..utils import save_image_bytes


def _parse_txt(data: bytes) -> ParsedBundle:
    text = data.decode(errors="ignore")
    return ParsedBundle(text=text, html=f"<pre>{BeautifulSoup(text, 'html.parser').text}</pre>")


def _parse_md(data: bytes) -> ParsedBundle:
    # Minimal: treat as text; client/editor can render further
    text = data.decode(errors="ignore")
    return ParsedBundle(text=text, html=None)


def _parse_pdf(data: bytes) -> ParsedBundle:
    reader = PdfReader(io.BytesIO(data))
    text_parts: List[str] = []
    for page in reader.pages:
        try:
            text_parts.append(page.extract_text() or "")
        except Exception:
            pass
    text = "\n\n".join(t.strip() for t in text_parts if t)
    # Images extraction from PDFs is non-trivial; recommend Docling for full fidelity
    return ParsedBundle(text=text)


def _parse_docx(data: bytes) -> ParsedBundle:
    # Extract text via python-docx
    buf = io.BytesIO(data)
    doc = DocxDocument(buf)
    paras = [p.text for p in doc.paragraphs]
    text = "\n".join(paras)

    # Extract embedded images by reading the docx as zip
    images: List[ImageRef] = []
    z = zipfile.ZipFile(io.BytesIO(data))
    for name in z.namelist():
        if name.startswith("word/media/"):
            img_bytes = z.read(name)
            url = save_image_bytes(img_bytes, Path(name).name)
            images.append(ImageRef(url=url))

    # Tables as plain text
    tables: List[TableRef] = []
    for t in doc.tables:
        cells = [[cell.text for cell in row.cells] for row in t.rows]
        tables.append(TableRef(data=cells))

    return ParsedBundle(text=text, images=images, tables=tables)


def _parse_html(data: bytes) -> ParsedBundle:
    soup = BeautifulSoup(data, "html.parser")
    text = soup.get_text("\n", strip=True)
    images: List[ImageRef] = []
    for img in soup.find_all("img"):
        src = img.get("src")
        if src and src.startswith("http"):
            images.append(ImageRef(url=src, alt=img.get("alt")))
    return ParsedBundle(text=text, html=str(soup), images=images)


def _parse_image_ocr(filename: str, data: bytes) -> ParsedBundle:
    # Save original image to static and run OCR over it
    try:
        im = Image.open(_io.BytesIO(data))
        # Convert to RGB to avoid mode issues
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        text = pytesseract.image_to_string(im)
    except Exception:
        text = ""
    # Persist image for later embedding
    url = save_image_bytes(data, Path(filename).name or "image.png")
    images = [ImageRef(url=url)]
    return ParsedBundle(text=text.strip(), images=images)


def _markdown_table(headers: List[str], rows: List[List[str]]) -> str:
    if not headers and rows:
        headers = [f"col{i+1}" for i in range(len(rows[0]))]
    if not headers:
        return ""
    sep = ["---" for _ in headers]
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(sep) + " |"]
    for r in rows:
        # pad/truncate to header length
        vals = [str(v) if v is not None else "" for v in r]
        if len(vals) < len(headers):
            vals = vals + [""] * (len(headers) - len(vals))
        else:
            vals = vals[: len(headers)]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def _parse_csv(data: bytes) -> ParsedBundle:
    try:
        text = data.decode("utf-8")
    except Exception:
        text = data.decode(errors="ignore")
    reader = _csv.reader(text.splitlines())
    rows = list(reader)
    headers: List[str] = rows[0] if rows else []
    body = rows[1:] if len(rows) > 1 else []
    md = _markdown_table(headers, body) if rows else ""
    return ParsedBundle(text=md or text)


def _parse_json(data: bytes) -> ParsedBundle:
    try:
        obj = _json.loads(data.decode("utf-8"))
    except Exception:
        try:
            obj = _json.loads(data.decode(errors="ignore"))
        except Exception:
            return ParsedBundle(text=data.decode(errors="ignore"))

    # If list of dicts -> table; else pretty JSON
    if isinstance(obj, list) and obj and isinstance(obj[0], dict):
        keys = list({k for item in obj if isinstance(item, dict) for k in item.keys()})
        rows = [[item.get(k, "") for k in keys] for item in obj if isinstance(item, dict)]
        md = _markdown_table(keys, rows)
        return ParsedBundle(text=md, tables=[TableRef(data=obj)])
    pretty = _json.dumps(obj, indent=2, ensure_ascii=False)
    return ParsedBundle(text=f"```json\n{pretty}\n```", tables=[TableRef(data=obj)])


def _parse_excel_xlsx(data: bytes) -> ParsedBundle:
    wb = load_workbook(_io.BytesIO(data), read_only=True, data_only=True)
    images: List[ImageRef] = []
    tables: List[TableRef] = []
    md_parts: List[str] = []
    for ws in wb.worksheets:
        rows: List[List[str]] = []
        for row in ws.iter_rows(values_only=True):
            rows.append(["" if v is None else str(v) for v in row])
        headers = rows[0] if rows else []
        body = rows[1:] if len(rows) > 1 else []
        md = _markdown_table(headers, body)
        if md:
            md_parts.append(f"### {ws.title}\n\n" + md)
        tables.append(TableRef(data=rows))
    text = "\n\n".join(md_parts)
    return ParsedBundle(text=text, tables=tables, images=images)


def parse_any(filename: str, data: bytes) -> ParsedBundle:
    suffix = Path(filename.lower()).suffix
    if suffix in {".txt"}:
        return _parse_txt(data)
    if suffix in {".md", ".markdown"}:
        return _parse_md(data)
    if suffix in {".csv"}:
        return _parse_csv(data)
    if suffix in {".json"}:
        return _parse_json(data)
    if suffix in {".pdf"}:
        return _parse_pdf(data)
    if suffix in {".docx"}:
        return _parse_docx(data)
    if suffix in {".html", ".htm"}:
        return _parse_html(data)
    if suffix in {".xlsx", ".xlsm"}:
        return _parse_excel_xlsx(data)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp", ".gif"}:
        return _parse_image_ocr(filename, data)
    # Fallback: best-effort text
    return _parse_txt(data)


def parse_with_docling(filename: str, data: bytes) -> Optional[ParsedBundle]:
    """Attempt to parse using Docling if available.

    Tries common Docling APIs and maps results to ParsedBundle.
    Falls back to None if Docling is not installed or an error occurs.
    """
    try:
        import io as _io
        import mimetypes

        try:
            # Preferred modern API
            from docling.document_converter import DocumentConverter, InputDocument  # type: ignore

            converter = DocumentConverter()
            mime, _ = mimetypes.guess_type(filename)
            doc = InputDocument(file_like=_io.BytesIO(data), filename=filename, mime_type=mime)
            res = converter.convert(doc)

            text = getattr(res, "text", None) or getattr(res, "plaintext", None)
            html = getattr(res, "html", None)
            images: List[ImageRef] = []
            if hasattr(res, "images") and res.images:
                for im in res.images:
                    name = getattr(im, "name", "image.png")
                    content = getattr(im, "content", None) or getattr(im, "bytes", None)
                    if content:
                        url = save_image_bytes(content, name)
                        images.append(ImageRef(url=url))
            tables: List[TableRef] = []
            if hasattr(res, "tables") and res.tables:
                for tbl in res.tables:
                    thtml = getattr(tbl, "html", None)
                    tdata = getattr(tbl, "data", None)
                    tables.append(TableRef(html=thtml, data=tdata))
            meta = getattr(res, "meta", None) or {}

            return ParsedBundle(text=text, html=html, images=images, tables=tables, meta=meta)
        except Exception:
            # Older or different API surface
            try:
                from docling import Document  # type: ignore
                d = Document.from_bytes(data, filename=filename)
                text = getattr(d, "text", None)
                html = getattr(d, "html", None)
                images: List[ImageRef] = []
                if hasattr(d, "images"):
                    for im in d.images:
                        content = getattr(im, "content", None) or getattr(im, "bytes", None)
                        name = getattr(im, "name", "image.png")
                        if content:
                            url = save_image_bytes(content, name)
                            images.append(ImageRef(url=url))
                tables: List[TableRef] = []
                if hasattr(d, "tables"):
                    for tbl in d.tables:
                        tables.append(TableRef(html=getattr(tbl, "html", None), data=getattr(tbl, "data", None)))
                return ParsedBundle(text=text, html=html, images=images, tables=tables)
            except Exception:
                return None
    except Exception:
        return None
